"""
01_build_fixtures.py

Generates the synthetic fixtures used in the benchmark:
  - complex.xlsx        (multi-sheet with merged cells, formulas, hidden sheet)
  - stress.xlsx         (14 sheets, each targeting a documented calamine limit)
  - realworld.xlsx      (messy real-world-style French quarterly report)
  - complex.html        (article page: nav/sidebar/tables/code/form/footer)
  - complex.docx        (technical design doc with headings, lists, tables)

Two PDFs used in the benchmark (weird_invoice.pdf, rust_book.pdf) are NOT
generated here because they are real-world inputs; they ship alongside
this script in /fixtures.

Run:  python3 01_build_fixtures.py
Deps: openpyxl, python-docx
"""
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.comments import Comment
from openpyxl.worksheet.datavalidation import DataValidation
from openpyxl.formatting.rule import CellIsRule, ColorScaleRule
from openpyxl.utils import get_column_letter
from openpyxl.workbook.defined_name import DefinedName
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date, datetime, time
import pathlib
import random

OUT = pathlib.Path(__file__).parent.parent / "fixtures"
OUT.mkdir(exist_ok=True)


# =========================================================================
# complex.xlsx — the "nice" test file (first iteration)
# =========================================================================
def build_complex_xlsx() -> None:
    random.seed(7)
    wb = Workbook()

    # Sheet 1: Q1 financials with merged regions and formulas
    ws = wb.active
    ws.title = "Q1 Financials"
    ws.merge_cells("A1:F1")
    ws["A1"] = "ACME Corp — Q1 2026 Financial Summary"
    ws["A1"].font = Font(bold=True, size=14, color="FFFFFF")
    ws["A1"].fill = PatternFill("solid", fgColor="2F4F4F")
    ws["A1"].alignment = Alignment(horizontal="center")
    ws.merge_cells("A2:F2")
    ws["A2"] = "All figures in USD thousands"
    ws["A2"].alignment = Alignment(horizontal="center")
    for i, h in enumerate(["Region", "Product", "Jan", "Feb", "Mar", "Q1 Total"], 1):
        c = ws.cell(row=4, column=i, value=h)
        c.font = Font(bold=True)
        c.fill = PatternFill("solid", fgColor="D3D3D3")
    regions = {
        "North America": ["Widget", "Gadget", "Gizmo"],
        "EMEA":          ["Widget", "Gadget"],
        "APAC":          ["Widget", "Gizmo", "Doohickey", "Thingamajig"],
    }
    r = 5
    for region, prods in regions.items():
        start = r
        for prod in prods:
            ws.cell(row=r, column=2, value=prod)
            for m in range(3):
                ws.cell(row=r, column=3 + m, value=random.randint(50, 500))
            ws.cell(row=r, column=6, value=f"=SUM(C{r}:E{r})")
            r += 1
        ws.merge_cells(start_row=start, start_column=1, end_row=r - 1, end_column=1)
        ws.cell(row=start, column=1, value=region).alignment = Alignment(vertical="center")
    ws.cell(row=r, column=1, value="GRAND TOTAL").font = Font(bold=True)
    ws.merge_cells(start_row=r, start_column=1, end_row=r, end_column=2)
    for col in range(3, 7):
        L = get_column_letter(col)
        ws.cell(row=r, column=col, value=f"=SUM({L}5:{L}{r - 1})").font = Font(bold=True)
    wb.defined_names["GrandTotal"] = DefinedName("GrandTotal", attr_text=f"'Q1 Financials'!$F${r}")

    # Sheet 2: Employees with a YEARFRAC formula
    ws2 = wb.create_sheet("Employees")
    ws2.append(["ID", "Name", "Department", "Hire Date", "Salary", "Active"])
    employees = [
        (1, "Alice Martin",    "Engineering", date(2019,  3, 14),  95000, True),
        (2, "Bob Chen",        "Engineering", date(2020, 11,  2),  87500, True),
        (3, "Carla Rossi",     "Sales",       date(2018,  7, 22),  72000, True),
        (4, "Daniel Okafor",   "Marketing",   date(2022,  1, 10),  68000, True),
        (5, "Eva Nguyen",      "Engineering", date(2023,  5, 30), 110000, True),
        (6, "François Dubois", "HR",          date(2015,  9,  1),  81000, False),
        (7, "Greta Svensson",  "Sales",       date(2021,  4, 15),  76500, True),
    ]
    for e in employees:
        ws2.append(e)
    ws2["G1"] = "Years"
    for i in range(2, len(employees) + 2):
        ws2[f"G{i}"] = f"=YEARFRAC(D{i},TODAY())"

    # Sheet 3: hidden
    ws3 = wb.create_sheet("Internal Notes")
    ws3["A1"] = "Confidential — do not share"
    ws3["A2"] = "Salary review due April 30"
    ws3.sheet_state = "hidden"

    # Sheet 4: numeric stress
    ws4 = wb.create_sheet("Matrix 20x20")
    for i in range(1, 21):
        for j in range(1, 21):
            ws4.cell(row=i, column=j, value=round(random.uniform(-1, 1), 4))

    wb.save(OUT / "complex.xlsx")
    print(f"  {OUT / 'complex.xlsx'}  → sheets: {wb.sheetnames}")


# =========================================================================
# stress.xlsx — one sheet per documented calamine limitation
# =========================================================================
def build_stress_xlsx() -> None:
    wb = Workbook()

    # Sheet 1 — number formats
    ws = wb.active
    ws.title = "1-Formats"
    ws.append(["Stored", "Format code", "What Excel displays (truth)"])
    rows = [
        (0.15,        "0%",                            "15%"),
        (0.1234,      "0.00%",                         "12.34%"),
        (1234.5,      "#,##0.00",                      "1,234.50"),
        (1234567,     "#,##0",                         "1,234,567"),
        (-1234.5,     "#,##0.00;[Red](#,##0.00)",      "(1,234.50)"),
        (0.000012345, "0.00E+00",                      "1.23E-05"),
        (7,           "000",                           "007"),
        (0.5,         "# ?/?",                         "1/2"),
        (3.14159265,  "0.00",                          "3.14"),
        (1234.56,     '"$"#,##0.00',                   "$1,234.56"),
        (0.99999,     "0.00%",                         "100.00%"),
    ]
    for stored, fmt, display in rows:
        r = ws.max_row + 1
        ws.cell(row=r, column=1, value=stored).number_format = fmt
        ws.cell(row=r, column=2, value=fmt)
        ws.cell(row=r, column=3, value=display)

    # Sheet 2 — dates, times, durations
    ws = wb.create_sheet("2-DatesTimes")
    ws.append(["Kind", "Stored value", "Format", "What Excel displays"])
    ws.append(["Date",     date(2026, 4, 21),             "yyyy-mm-dd",       "2026-04-21"])
    ws.append(["Datetime", datetime(2026, 4, 21, 14, 30), "yyyy-mm-dd hh:mm", "2026-04-21 14:30"])
    ws.append(["Time",     time(9, 45, 15),               "hh:mm:ss",         "09:45:15"])
    r = ws.max_row + 1
    ws.cell(row=r, column=1, value="Duration 49h")
    ws.cell(row=r, column=2, value=49 / 24 + 30 / 1440).number_format = "[h]:mm"
    ws.cell(row=r, column=3, value="[h]:mm")
    ws.cell(row=r, column=4, value="49:30")
    ws.append(["Old date", date(1900, 3, 1), "yyyy-mm-dd", "1900-03-01"])

    # Sheet 3 — merged cells
    ws = wb.create_sheet("3-Merged")
    ws["A1"] = "Region"; ws["B1"] = "Product"; ws["C1"] = "Revenue"
    ws["A2"] = "North America"; ws["B2"] = "Widget"; ws["C2"] = 1000
    ws["B3"] = "Gadget";  ws["C3"] = 2000
    ws["B4"] = "Gizmo";   ws["C4"] = 3000
    ws.merge_cells("A2:A4")
    ws["A5"] = "EMEA"; ws["B5"] = "Widget"; ws["C5"] = 500
    ws["B6"] = "Gadget"; ws["C6"] = 800
    ws.merge_cells("A5:A6")
    ws.merge_cells("A8:C8")
    ws["A8"] = "GRAND TOTAL"
    ws["A8"].font = Font(bold=True)
    ws["A8"].alignment = Alignment(horizontal="center")
    ws.merge_cells("E1:F2")
    ws["E1"] = "2x2 Header"
    ws["E1"].alignment = Alignment(horizontal="center", vertical="center")

    # Sheet 4 — hyperlinks
    ws = wb.create_sheet("4-Links")
    ws.append(["Label", "Target"])
    ws["A2"] = "Our website"
    ws["A2"].hyperlink = "https://www.example.com"
    ws["A2"].font = Font(color="0000FF", underline="single")
    ws["A3"] = "Internal ref"
    ws["A3"].hyperlink = "#Sheet1!A1"
    ws["A4"] = "Email us"
    ws["A4"].hyperlink = "mailto:hello@example.com"
    ws["B2"] = "https://www.example.com"
    ws["B3"] = "#Sheet1!A1"
    ws["B4"] = "mailto:hello@example.com"

    # Sheet 5 — comments
    ws = wb.create_sheet("5-Comments")
    ws["A1"] = "Revenue Q1"
    ws["A1"].comment = Comment("Includes the Marlow-Chen deal (est. $2M)", "Auditor")
    ws["A2"] = "Headcount"
    ws["A2"].comment = Comment("42 confirmed, 3 pending offers", "HR")
    ws["B1"] = 12500000
    ws["B2"] = 42

    # Sheet 6 — styles
    ws = wb.create_sheet("6-Styles")
    ws["A1"] = "Red text"
    ws["A1"].font = Font(color="FF0000", bold=True)
    ws["B1"] = "Highlighted"
    ws["B1"].fill = PatternFill("solid", fgColor="FFFF00")
    ws["A2"] = "Bordered"
    ws["A2"].border = Border(
        left=Side(style="thick"), right=Side(style="thick"),
        top=Side(style="thick"), bottom=Side(style="thick"),
    )

    # Sheet 7 — line breaks
    ws = wb.create_sheet("7-LineBreaks")
    ws.append(["Kind", "Value"])
    ws["A2"] = "LF in cell"
    ws["B2"] = "line1\nline2\nline3"
    ws["B2"].alignment = Alignment(wrap_text=True)
    ws["A3"] = "Text with \\r\\n"
    ws["B3"] = "before\r\nafter"

    # Sheet 8 — formulas & errors
    ws = wb.create_sheet("8-Formulas")
    ws.append(["Formula", "Expected value"])
    ws["A2"] = "=1/0";                            ws["B2"] = "#DIV/0!"
    ws["A3"] = '=VLOOKUP("x",A1:B1,2,FALSE)';    ws["B3"] = "#N/A"
    ws["A4"] = "=SUM(B2:B3)";                     ws["B4"] = "Depends on above"
    ws["A5"] = "=A1";                             ws["B5"] = "A1 of this sheet"
    ws["A6"] = "=SQRT(-1)";                       ws["B6"] = "#NUM!"

    # Sheet 9 — data validation
    ws = wb.create_sheet("9-Validation")
    ws.append(["Status", "Score"])
    ws.append(["Open", 5])
    ws.append(["Closed", 10])
    dv = DataValidation(type="list", formula1='"Open,Pending,Closed"', allow_blank=True)
    dv.add("A2:A10")
    ws.add_data_validation(dv)
    dv2 = DataValidation(type="whole", operator="between", formula1=1, formula2=10)
    dv2.add("B2:B10")
    ws.add_data_validation(dv2)

    # Sheet 10 — conditional formatting
    ws = wb.create_sheet("10-Conditional")
    for i, v in enumerate([5, 10, 15, 20, 25, 30], start=1):
        ws.cell(row=i, column=1, value=v)
    ws.conditional_formatting.add(
        "A1:A6",
        CellIsRule(operator="greaterThan", formula=["20"],
                   fill=PatternFill("solid", fgColor="FF9999")),
    )

    # Sheet 11 — hidden rows/cols
    ws = wb.create_sheet("11-Hidden")
    for i, v in enumerate(["visible1", "HIDDEN_ROW_SECRET", "visible2", "visible3"], start=1):
        ws.cell(row=i, column=1, value=v)
    ws.row_dimensions[2].hidden = True
    ws["A1"] = "visible1"; ws["B1"] = "also visible"; ws["C1"] = "HIDDEN_COL_SECRET"
    ws.column_dimensions["C"].hidden = True

    # Sheet 12 — very hidden sheet
    ws = wb.create_sheet("12-VeryHidden")
    ws["A1"] = "VERY_HIDDEN_CONTENT — should normally not surface in UI"
    ws.sheet_state = "veryHidden"

    # Sheet 13 — named ranges
    ws = wb.create_sheet("13-NamedRanges")
    ws["A1"] = 100; ws["A2"] = 200; ws["A3"] = 300
    ws["B1"] = "=SUM(MyTotal)"
    wb.defined_names["MyTotal"] = DefinedName("MyTotal", attr_text="'13-NamedRanges'!$A$1:$A$3")
    wb.defined_names["TaxRate"] = DefinedName("TaxRate", attr_text="0.2")

    # Sheet 14 — messy header
    ws = wb.create_sheet("14-MessyHeader")
    ws.merge_cells("A2:D2")
    ws["A2"] = "QUARTERLY REPORT — 2026"
    ws["A2"].alignment = Alignment(horizontal="center")
    ws["A2"].font = Font(bold=True, size=14)
    ws["A4"] = "Metric"; ws["B4"] = "Jan"; ws["C4"] = "Feb"; ws["D4"] = "Mar"
    ws["A5"] = "Revenue"; ws["B5"] = 1000; ws["C5"] = 1200; ws["D5"] = 1100
    ws["A6"] = "Costs";   ws["B6"] = 800;  ws["C6"] = 900;  ws["D6"] = 950

    wb.save(OUT / "stress.xlsx")
    print(f"  {OUT / 'stress.xlsx'}  → {len(wb.sheetnames)} sheets")


# =========================================================================
# realworld.xlsx — messy French quarterly report
# =========================================================================
def build_realworld_xlsx() -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "data"

    ws.merge_cells("A1:H1")
    ws["A1"] = "RAPPORT TRIMESTRIEL - CONFIDENTIEL"
    ws["A1"].font = Font(bold=True, size=16, color="FFFFFF")
    ws["A1"].fill = PatternFill("solid", fgColor="1F4E79")
    ws["A1"].alignment = Alignment(horizontal="center")
    ws.merge_cells("A2:H2")
    ws["A2"] = "Préparé par: Jean-François Müller  |  Date: 2026-04-15  |  Version: 2.3 (définitive)"
    ws["A2"].alignment = Alignment(horizontal="center")

    ws.merge_cells("A4:B7")
    ws["A4"] = "[LOGO]"
    ws["A4"].alignment = Alignment(horizontal="center", vertical="center")
    ws["A4"].font = Font(italic=True, color="999999")

    ws["C4"] = "Client:";     ws["D4"] = "ACME Industries SARL"
    ws["C5"] = "Contrat n°:"; ws["D5"] = "2026-047-A"
    ws["C6"] = "Montant HT:"; ws["D6"] = 145750.50; ws["D6"].number_format = '#,##0.00 "€"'
    ws["C7"] = "TVA (20%):";  ws["D7"] = 29150.10;  ws["D7"].number_format = '#,##0.00 "€"'
    ws["E4"] = "Début:";      ws["F4"] = date(2026, 1, 15)
    ws["E5"] = "Fin prévue:"; ws["F5"] = date(2026, 6, 30)
    ws["E6"] = "Statut:";     ws["F6"] = "En cours"
    ws["F6"].fill = PatternFill("solid", fgColor="FFC000")

    ws.merge_cells("A9:A10")
    ws["A9"] = "Projet"
    ws.merge_cells("B9:C9"); ws["B9"] = "Budget"; ws["B10"] = "Prévu"; ws["C10"] = "Réel"
    ws.merge_cells("D9:E9"); ws["D9"] = "Délai";  ws["D10"] = "Début"; ws["E10"] = "Fin"
    ws.merge_cells("F9:F10"); ws["F9"] = "Responsable"
    ws.merge_cells("G9:G10"); ws["G9"] = "% avanc."
    ws.merge_cells("H9:H10"); ws["H9"] = "Notes"
    for cell in ["A9", "B9", "D9", "F9", "G9", "H9", "B10", "C10", "D10", "E10"]:
        ws[cell].font = Font(bold=True)
        ws[cell].fill = PatternFill("solid", fgColor="D9E1F2")
        ws[cell].alignment = Alignment(horizontal="center")

    projects = [
        ("Migration ERP",    45000, 48200, date(2026, 1, 20), date(2026, 4, 30), "A. Dupont",   0.85, "Retard de 2 sem. dû à intégration SSO"),
        ("Refonte site web", 18500, 16200, date(2026, 2, 1),  date(2026, 3, 15), "M. Chen",     1.0,  ""),
        ("Audit sécurité",   25000, 25000, date(2026, 3, 1),  date(2026, 3, 31), "Y. Saïd",     1.0,  "RAS"),
        ("Formation DPO",    8000,  None,  date(2026, 4, 15), date(2026, 5, 30), "C. Álvarez",  0.10, "Démarrage cette semaine"),
        ("POC IA générative", 30000, 42100, date(2026, 2, 15), date(2026, 6, 30), "J-F. Müller", 0.60, "Budget dépassé — cf. note #W-023"),
        ("", None, None, None, None, "", None, ""),
        ("Total", "=SUM(B11:B15)", "=SUM(C11:C15)", None, None, None, None, None),
    ]
    for i, p in enumerate(projects, start=11):
        for j, val in enumerate(p, start=1):
            c = ws.cell(row=i, column=j, value=val)
            if j in (2, 3) and isinstance(val, (int, float)):
                c.number_format = '#,##0 "€"'
            if j == 7 and isinstance(val, float):
                c.number_format = "0%"
            if j in (4, 5) and val:
                c.number_format = "dd/mm/yyyy"
    for col in "ABCDEFGH":
        ws[f"{col}17"].font = Font(bold=True)
        ws[f"{col}17"].fill = PatternFill("solid", fgColor="FFE699")

    ws["C15"].comment = Comment(
        "Budget dépassé de 40% — escaladé à la direction le 2026-03-28",
        "audit@acme",
    )
    ws.conditional_formatting.add(
        "G11:G15",
        ColorScaleRule(start_type="num", start_value=0, start_color="F8696B",
                       mid_type="num", mid_value=0.5, mid_color="FFEB84",
                       end_type="num", end_value=1, end_color="63BE7B"),
    )

    ws2 = wb.create_sheet("Notes")
    ws2["A1"] = "Décisions prises:"
    ws2["A2"] = "1. Budget IA réévalué de 30k à 45k (décision COMEX 2026-03-10)"
    ws2["A3"] = "2. Migration ERP: bascule prod reportée au 2026-05-15"
    ws2["A5"] = "À faire:"
    ws2["A6"] = "- Valider devis fournisseur SSO"
    ws2["A7"] = "- Planifier réunion post-mortem audit"

    wb.save(OUT / "realworld.xlsx")
    print(f"  {OUT / 'realworld.xlsx'}  → {len(wb.sheetnames)} sheets")


# =========================================================================
# complex.html
# =========================================================================
def build_complex_html() -> None:
    html = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8">
<title>Deep Dive: Rust Ownership — TechBlog</title>
<meta name="description" content="An illustrated walkthrough of Rust's ownership model.">
<style>.nav{background:#222;color:#fff}.sidebar{float:right;width:30%}</style>
</head>
<body>
<nav class="nav">
  <ul><li><a href="/">Home</a></li><li><a href="/archive">Archive</a></li><li><a href="/about">About</a></li></ul>
</nav>

<aside class="sidebar">
  <h3>Related</h3>
  <ul>
    <li><a href="/posts/lifetimes">Understanding Lifetimes</a></li>
    <li><a href="/posts/borrow-checker">The Borrow Checker Explained</a></li>
  </ul>
  <div class="ad">
    <p>Sponsored: <a href="https://example.com">Get the Rust Cookbook</a></p>
  </div>
</aside>

<main>
<article>
<h1>Deep Dive: Rust Ownership</h1>
<p class="byline">By <a href="/author/jane">Jane Doe</a> · Published 12 March 2026 · 8&nbsp;min read</p>

<p>Rust's ownership model is what makes it <em>unique</em> among systems languages. In this article we'll
walk through the three rules, with examples.</p>

<h2>The Three Rules</h2>
<ol>
  <li>Each value has a single <strong>owner</strong>.</li>
  <li>When the owner goes out of scope, the value is <strong>dropped</strong>.</li>
  <li>You can have <em>either</em>
    <ul>
      <li>one mutable reference, <strong>or</strong></li>
      <li>any number of immutable references</li>
    </ul>
    — but never both at the same time.</li>
</ol>

<h2>A First Example</h2>
<p>Consider this code:</p>
<pre><code class="language-rust">fn main() {
    let s = String::from("hello");
    takes_ownership(s);
}

fn takes_ownership(s: String) {
    println!("{}", s);
}</code></pre>

<h2>Comparison Table</h2>
<table>
  <thead>
    <tr><th>Operation</th><th>Stack</th><th>Heap</th><th>Notes</th></tr>
  </thead>
  <tbody>
    <tr><td>Copy</td><td>implicit (<code>Copy</code>)</td><td>requires <code>.clone()</code></td><td>primitive types only on stack</td></tr>
    <tr><td>Move</td><td>—</td><td>invalidates source</td><td>default for <code>String</code>, <code>Vec</code>…</td></tr>
    <tr><td>Borrow</td><td>cheap</td><td>cheap</td><td>checked at compile time</td></tr>
  </tbody>
</table>

<blockquote>
  <p>"Ownership is Rust's most unique feature." — <cite>The Rust Book</cite></p>
</blockquote>

<h2>Leave a Comment</h2>
<form action="/comment" method="post">
  <label>Name: <input type="text" name="name" required></label>
  <label>Comment: <textarea name="body"></textarea></label>
  <button type="submit">Post</button>
</form>
</article>
</main>

<footer>
  <p>&copy; 2026 TechBlog. <a href="/privacy">Privacy</a> · <a href="/tos">Terms</a></p>
  <script>console.log("tracking")</script>
</footer>
</body>
</html>
"""
    (OUT / "complex.html").write_text(html, encoding="utf-8")
    print(f"  {OUT / 'complex.html'}")


# =========================================================================
# complex.docx
# =========================================================================
def build_complex_docx() -> None:
    doc = Document()

    t = doc.add_heading("Technical Design Document", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Project: PaymentGateway v2  ·  Author: J. Martinez  ·  Draft 0.3 — 2026-04-10"
                      ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("1. Executive Summary", 1)
    p = doc.add_paragraph("This document describes the ")
    p.add_run("architecture").bold = True
    p.add_run(" and ")
    p.add_run("trade-offs").italic = True
    p.add_run(" of the next-generation payment gateway.")

    doc.add_heading("2. Requirements", 1)
    doc.add_heading("2.1 Functional", 2)
    for item in ["Process 5,000 TPS peak", "Support EUR, USD, GBP, CHF", "Sub-200ms p99 latency"]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("2.2 Non-functional", 2)
    for item in ["99.99% uptime", "PCI-DSS compliance", "Multi-region active-active"]:
        doc.add_paragraph(item, style="List Number")

    doc.add_heading("4. Service Inventory", 1)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    hdr[0].text = "Service"; hdr[1].text = "Owner"; hdr[2].text = "SLO"; hdr[3].text = "Datastore"
    for r in [
        ("auth-svc",         "Identity team", "99.99%",  "Postgres 15"),
        ("ledger-svc",       "Ledger team",   "99.999%", "CockroachDB"),
        ("fx-svc",           "Pricing team",  "99.95%",  "Redis + Kafka"),
        ("notification-svc", "Platform team", "99.9%",   "DynamoDB"),
    ]:
        row = tbl.add_row().cells
        for i, v in enumerate(r):
            row[i].text = v

    doc.add_heading("5. Risks & Mitigations", 1)
    p = doc.add_paragraph()
    p.add_run("⚠ ").font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    p.add_run("The PCI-DSS audit must complete before general availability.")

    doc.add_heading("6. Sample Config", 1)
    p = doc.add_paragraph()
    run = p.add_run("service:\n  name: ledger-svc\n  replicas: 12\n  resources:\n    cpu: 4\n    memory: 8Gi")
    run.font.name = "Courier New"
    run.font.size = Pt(10)

    doc.add_paragraph("See also: https://wiki.acme.example/payments/v2 for diagrams.")
    doc.save(OUT / "complex.docx")
    print(f"  {OUT / 'complex.docx'}")


if __name__ == "__main__":
    print("Building fixtures in", OUT)
    build_complex_xlsx()
    build_stress_xlsx()
    build_realworld_xlsx()
    build_complex_html()
    build_complex_docx()
    print("Done.")
